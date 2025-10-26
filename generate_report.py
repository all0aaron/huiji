# -*- coding: utf-8 -*-
"""
安井食品并购新宏业开题报告Word文档生成器
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_heading(doc, text, level=1):
    """创建标题并设置中文字体"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        if level == 1:
            run.font.size = Pt(16)
            run.font.bold = True
        elif level == 2:
            run.font.size = Pt(14)
            run.font.bold = True
    return heading

def add_paragraph(doc, text, font_size=12):
    """添加段落并设置格式"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.first_line_indent = Inches(0.35)
    paragraph.paragraph_format.line_spacing = 1.5
    
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(font_size)
    
    return paragraph

print("开始生成Word文档...")

# 创建Word文档
doc = Document()

# 设置默认字体
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
doc.styles['Normal'].font.size = Pt(12)

# 封面
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_before = Pt(72)

run = title.add_run('本科生毕业论文（设计）开题报告')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run.font.size = Pt(22)
run.font.bold = True

doc.add_paragraph('\n\n')

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('题目：安井食品并购新宏业动因及绩效分析')
run.font.name = 'Times New Roman'
run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run.font.size = Pt(18)
run.font.bold = True

doc.add_paragraph('\n\n\n')

# 学生信息表格
table = doc.add_table(rows=7, cols=2)
table.alignment = WD_ALIGN_PARAGRAPH.CENTER

info_data = [
    ('学院', '经济管理学院'),
    ('专业', '会计学'),
    ('班级', '22本会计2班'),
    ('学号', '202299008892'),
    ('姓名', '吴宇萍'),
    ('指导教师1', '官晓风'),
    ('指导教师2', '')
]

for i, (label, value) in enumerate(info_data):
    row = table.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    
    for cell in row.cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            run.font.size = Pt(12)

doc.add_page_break()

# 第一章：绪论
create_heading(doc, '第1章 绪论', level=1)

create_heading(doc, '1.1 研究背景与研究意义', level=2)

create_heading(doc, '1.1.1 研究背景', level=3)
add_paragraph(doc, '在全球经济一体化和市场竞争日益激烈的背景下，企业并购已成为企业实现快速扩张、优化资源配置、提升市场竞争力的重要战略手段。近年来，我国食品行业并购重组活动频繁，企业通过并购实现产业链整合、市场份额扩大和品牌价值提升。安井食品股份有限公司作为速冻食品行业的领军企业，自2001年成立以来，始终专注于速冻火锅料制品和速冻面米制品的研发、生产与销售。')

add_paragraph(doc, '2021年，安井食品通过并购新宏业食品有限公司，进一步完善了其产品线布局，拓展了市场渠道，增强了企业的综合竞争力。新宏业作为区域性知名速冻食品企业，在华东地区拥有稳定的客户群体和成熟的销售网络。此次并购不仅有助于安井食品实现规模经济效应，还能通过资源整合和协同效应提升企业的盈利能力和市场地位。')

print("正在保存Word文档...")

# 保存文档
filename = '安井食品并购新宏业开题报告_吴宇萍.docx'
doc.save(filename)

print(f"\n✅ Word文档生成成功！")
print(f"文件名：{filename}")
print("请查看当前目录")

